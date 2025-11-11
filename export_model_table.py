# export_model_table.py
import math
import re
import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_model_table(model_result, file_path):
    def prettify_terms(term):
        term = term.replace(":", " * ")
        if term in ("Intercept", "const", "(Intercept)", "Intercept[0]"):
            return "Intercept"
        term = re.sub(r'([A-Za-z_][A-Za-z0-9_]*)(\[T\.([^\]]+)\])', r'\1 [\3]', term)
        term = re.sub(r'np\.power\(([^,]+),\s*2\)', r'\1^2', term)
        return term

    def p_fmt(p):
        if p is None or (isinstance(p, float) and (math.isnan(p) or math.isinf(p))):
            return ""
        if p <= 0.001:
            return f"{p:.2e}"
        return f"{p:.3f}"

    def p_stars(p):
        if p is None or (isinstance(p, float) and (math.isnan(p) or math.isinf(p))):
            return ""
        return "***" if p < 0.001 else ("**" if p < 0.01 else ("*" if p < 0.05 else ""))

    def get_title(res):
        # prefer formula when available
        try:
            f = getattr(res.model, "formula", None)
            if isinstance(f, str):
                return f
        except Exception:
            pass
        try:
            y = getattr(res.model, "endog_names", "y")
            X = getattr(res.model, "exog_names", [])
            if isinstance(X, (list, tuple)):
                rhs = " + ".join([x for x in X if x not in ("Intercept", "const")])
            else:
                rhs = str(X)
            return f"{y} ~ {rhs}" if rhs else y
        except Exception:
            return "Model"

    def as_series(obj):
        return None if obj is None else pd.Series(obj, index=params.index, dtype=float)

    def confint_df(res):
        try:
            ci = res.conf_int()
            if isinstance(ci, pd.DataFrame):
                ci.columns = ["ci_low", "ci_high"]
                return ci
            ci = pd.DataFrame(ci, index=params.index, columns=["ci_low", "ci_high"])
            return ci
        except Exception:
            if seS is None:
                return None
            lo = params - 1.96 * seS
            hi = params + 1.96 * seS
            return pd.DataFrame({"ci_low": lo, "ci_high": hi})

    def safe_round(x, k):
        try:
            return np.round(float(x), k)
        except Exception:
            return np.nan

    # ---------- fixed effects ----------
    params = getattr(model_result, "params", None)
    if params is None:
        raise ValueError("Provided result has no .params; pass a fitted statsmodels result.")

    if isinstance(params, (pd.Series, pd.DataFrame)):
        params = params.squeeze()
    params = pd.Series(params)

    # Standard errors: try .bse, then .bse_fe (MixedLM)
    seS = getattr(model_result, "bse", None)
    if seS is None and hasattr(model_result, "bse_fe"):
        seS = getattr(model_result, "bse_fe")
        # align to fixed-effect names if present
        try:
            fe_names = getattr(model_result.model, "exog_names", None)
            if isinstance(fe_names, list):
                seS = pd.Series(seS, index=fe_names)
        except Exception:
            pass
    seS = as_series(seS)

    # t/z statistics
    statS = as_series(getattr(model_result, "tvalues", None))
    if statS is None:
        statS = as_series(getattr(model_result, "zvalues", None))

    # p-values
    pS = as_series(getattr(model_result, "pvalues", None))

    # If anything missing, compute Wald stats/p from params & SE
    if statS is None and seS is not None:
        statS = (params / seS)
    if pS is None and statS is not None:
        # two-sided normal approximation (matches statsmodels MixedLM)
        pS = 2.0 * (1.0 - stats.norm.cdf(np.abs(statS)))

    ci = confint_df(model_result)

    df = pd.DataFrame({
        "Variable": [prettify_terms(t) for t in params.index],
        "β": [safe_round(v, 3) for v in params.values],
        "SE": [safe_round(seS.get(i), 3) if seS is not None and i in seS.index else np.nan for i in params.index],
        "t/z-value": [safe_round(statS.get(i), 3) if statS is not None and i in statS.index else np.nan for i in params.index],
        "p_raw": [pS.get(i) if pS is not None and i in pS.index else np.nan for i in params.index]
    })
    if ci is not None:
        ci = ci.reindex(params.index)
        df["CI"] = [f"{safe_round(lo,3)} – {safe_round(hi,3)}" for lo, hi in zip(ci["ci_low"], ci["ci_high"])]
    else:
        df["CI"] = ""

    df["p-value"] = [p_fmt(p) for p in df["p_raw"]]
    df["stars"] = [p_stars(p) for p in df["p_raw"]]
    fixed_tbl = df[["Variable", "β", "SE", "CI", "t/z-value", "p-value", "stars"]]

    # ---------- variance components ----------
    var_rows = []

    # Mixed effects: SDs from random-effects covariance
    cov_re = getattr(model_result, "cov_re", None)
    if cov_re is not None:
        try:
            sds = np.sqrt(np.diag(np.asarray(cov_re)))
            names = getattr(getattr(model_result, "model", None), "exog_re_names", None)
            names = list(names) if names is not None else [f"RE_{i+1}" for i in range(len(sds))]
            for nm, sd in zip(names, sds):
                label = "Random intercept" if nm in ("Intercept", "(Intercept)") else nm
                var_rows.append({"Variance": label, "SD": np.round(sd, 2)})
        except Exception:
            pass

    # Residual SD
    scale = getattr(model_result, "scale", None)
    if scale is None:
        scale = getattr(model_result, "mse_resid", None)
    if scale is not None and not (isinstance(scale, float) and math.isnan(scale)):
        var_rows.append({"Variance": "Residual", "SD": np.round(math.sqrt(float(scale)), 2)})

    if not var_rows:
        var_rows = [{"Variance": "Residual", "SD": np.nan}]

    var_tbl = pd.DataFrame(var_rows, columns=["Variance", "SD"])
    ll = getattr(model_result, "llf", np.nan)
    var_tbl["Goodness of fit"] = ""
    var_tbl.loc[var_tbl.index[0], "Goodness of fit"] = f"Log likelihood  {np.round(ll, 1) if not (isinstance(ll, float) and math.isnan(ll)) else ''}"

    # ---------- Word document ----------
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    title = get_title(model_result)
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    doc.add_paragraph("")

    headers = ["Variable", "β", "SE", "CI", "t-value" if "t" in fixed_tbl["t/z-value"].astype(str).to_string() else "t/z-value", "p-value"]
    table = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        run = table.rows[0].cells[j].paragraphs[0].add_run(h); run.bold = True

    for _, row in fixed_tbl.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Variable"])
        cells[1].text = "" if np.isnan(row["β"]) else f"{row['β']:.3f}"
        cells[2].text = "" if np.isnan(row["SE"]) else f"{row['SE']:.3f}"
        cells[3].text = str(row["CI"])
        cells[4].text = "" if np.isnan(row["t/z-value"]) else f"{row['t/z-value']:.3f}"
        pcell = cells[5].paragraphs[0]
        pcell.add_run(row["p-value"] if isinstance(row["p-value"], str) else "")
        if row["stars"]:
            star = pcell.add_run(row["stars"]); star.bold = True

    doc.add_paragraph("")
    t = doc.add_paragraph(); rr = t.add_run("Variance components"); rr.bold = True; rr.font.size = Pt(12)
    doc.add_paragraph("")
    vheaders = ["Variance", "SD", "Goodness of fit"]
    vtable = doc.add_table(rows=1, cols=len(vheaders))
    for j, h in enumerate(vheaders):
        run = vtable.rows[0].cells[j].paragraphs[0].add_run(h); run.bold = True
    for _, row in var_tbl.iterrows():
        cells = vtable.add_row().cells
        cells[0].text = str(row["Variance"])
        cells[1].text = "" if pd.isna(row["SD"]) else f"{row['SD']:.2f}"
        cells[2].text = str(row["Goodness of fit"])

    for tbl in (table, vtable):
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(file_path)
    